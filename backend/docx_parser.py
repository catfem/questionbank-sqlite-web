import re
import json
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import logging

logger = logging.getLogger(__name__)

class ParsedQuestion:
    def __init__(self):
        self.stem = ""
        self.question_type = None
        self.options = []
        self.correct_answer = []
        self.explanation = ""
        self.difficulty = "medium"
        self.tags = []
        self.raw_lines = []

class DocxParser:
    def __init__(self):
        self.questions = []
        self.errors = []
        self.current_question = None
        
    def parse_document(self, file_path: str) -> Tuple[List[ParsedQuestion], List[Dict]]:
        """Parse a DOCX file and extract questions."""
        try:
            doc = Document(file_path)
            self._process_paragraphs(doc.paragraphs)
            self._finalize_current_question()
            return self.questions, self.errors
        except Exception as e:
            logger.error(f"Error parsing document: {e}")
            raise e
    
    def _process_paragraphs(self, paragraphs):
        """Process all paragraphs in the document."""
        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Try to detect question patterns
            if self._is_question_start(text):
                if self.current_question:
                    self._finalize_current_question()
                self.current_question = ParsedQuestion()
                self.current_question.raw_lines.append((i+1, text))
                self._parse_question_stem(text)
            elif self.current_question:
                self.current_question.raw_lines.append((i+1, text))
                self._parse_question_content(text)
    
    def _is_question_start(self, text: str) -> bool:
        """Check if text starts a new question."""
        patterns = [
            r'^\d+\.\s+',  # "1. ", "2. ", etc.
            r'^\d+\)\s+',  # "1) ", "2) ", etc.
            r'^Question\s+\d+',  # "Question 1", "Question 2", etc.
            r'^Q\d+\.\s+',  # "Q1. ", "Q2. ", etc.
        ]
        return any(re.match(pattern, text, re.IGNORECASE) for pattern in patterns)
    
    def _parse_question_stem(self, text: str):
        """Parse the question stem."""
        # Remove question number
        stem = re.sub(r'^(\d+\.|\d+\)|Question\s+\d+|Q\d+\.)\s*', '', text, flags=re.IGNORECASE)
        self.current_question.stem = stem.strip()
    
    def _parse_question_content(self, text: str):
        """Parse options, answers, and explanations."""
        # Check if it's an option
        option_match = re.match(r'^([A-Z])[\.\)]\s*(.+)', text, re.IGNORECASE)
        if option_match:
            label = option_match.group(1).upper()
            option_text = option_match.group(2).strip()
            self.current_question.options.append({
                'label': label,
                'text': option_text
            })
            return
        
        # Check if it's an answer line
        answer_match = re.match(r'^(Answer|Correct|Solution)[:\s]+(.+)', text, re.IGNORECASE)
        if answer_match:
            answer_text = answer_match.group(2).strip()
            self._parse_answer(answer_text)
            return
        
        # Check if it's an explanation
        explanation_match = re.match(r'^(Explanation|Reasoning|Explain)[:\s]+(.+)', text, re.IGNORECASE)
        if explanation_match:
            self.current_question.explanation = explanation_match.group(2).strip()
            return
        
        # Check if it's a difficulty indicator
        difficulty_match = re.match(r'^(Difficulty|Level)[:\s]+(.+)', text, re.IGNORECASE)
        if difficulty_match:
            self.current_question.difficulty = difficulty_match.group(2).strip().lower()
            return
        
        # If none of the above, append to stem (multi-line question)
        if self.current_question.stem and not self.current_question.options:
            self.current_question.stem += " " + text
    
    def _parse_answer(self, answer_text: str):
        """Parse the answer text to determine correct options."""
        # True/False answers
        if answer_text.upper() in ['TRUE', 'FALSE', 'T', 'F']:
            self.current_question.question_type = 'true_false'
            if answer_text.upper() in ['TRUE', 'T']:
                self.current_question.correct_answer = [0]
            else:
                self.current_question.correct_answer = [1]
            # Add True/False options if not present
            if not self.current_question.options:
                self.current_question.options = [
                    {'label': 'A', 'text': 'True'},
                    {'label': 'B', 'text': 'False'}
                ]
            return
        
        # Handle various answer formats
        # Single letter: "A", "B", etc.
        if re.match(r'^[A-Z]$', answer_text.upper()):
            self.current_question.correct_answer = [ord(answer_text.upper()) - ord('A')]
            self._determine_question_type()
            return
        
        # Multiple letters: "A, C", "A and C", etc.
        letters = re.findall(r'[A-Z]', answer_text.upper())
        if letters:
            self.current_question.correct_answer = [ord(letter) - ord('A') for letter in letters]
            self._determine_question_type()
            return
    
    def _determine_question_type(self):
        """Determine question type based on number of correct answers."""
        if not self.current_question.question_type:
            if len(self.current_question.correct_answer) == 1:
                self.current_question.question_type = 'single'
            else:
                self.current_question.question_type = 'multiple'
    
    def _finalize_current_question(self):
        """Finalize the current question and add to the list."""
        if not self.current_question:
            return
        
        # Validation checks
        if not self.current_question.stem:
            self.errors.append({
                'line_number': self.current_question.raw_lines[0][0] if self.current_question.raw_lines else 0,
                'content': ' '.join([line[1] for line in self.current_question.raw_lines]),
                'error': 'Missing question stem'
            })
            self.current_question = None
            return
        
        # Auto-create True/False options if type is true_false but no options exist
        if self.current_question.question_type == 'true_false' and not self.current_question.options:
            self.current_question.options = [
                {'label': 'A', 'text': 'True'},
                {'label': 'B', 'text': 'False'}
            ]
        
        if not self.current_question.options and self.current_question.question_type != 'true_false':
            self.errors.append({
                'line_number': self.current_question.raw_lines[0][0] if self.current_question.raw_lines else 0,
                'content': ' '.join([line[1] for line in self.current_question.raw_lines]),
                'error': 'Missing options for non-true/false question'
            })
            self.current_question = None
            return
        
        if not self.current_question.correct_answer:
            self.errors.append({
                'line_number': self.current_question.raw_lines[0][0] if self.current_question.raw_lines else 0,
                'content': ' '.join([line[1] for line in self.current_question.raw_lines]),
                'error': 'Missing correct answer'
            })
            self.current_question = None
            return
        
        # Validate correct answer indices
        max_index = len(self.current_question.options) - 1
        for answer_idx in self.current_question.correct_answer:
            if answer_idx > max_index:
                self.errors.append({
                    'line_number': self.current_question.raw_lines[0][0] if self.current_question.raw_lines else 0,
                    'content': ' '.join([line[1] for line in self.current_question.raw_lines]),
                    'error': f'Correct answer index {answer_idx} exceeds options count'
                })
                self.current_question = None
                return
        
        # Set default question type if not determined
        if not self.current_question.question_type:
            self.current_question.question_type = 'single'
        
        self.questions.append(self.current_question)
        self.current_question = None

def create_sample_docx(file_path: str, questions_data: List[Dict]):
    """Create a sample DOCX file with questions for testing."""
    doc = Document()
    
    for i, q_data in enumerate(questions_data, 1):
        # Question number and stem
        p = doc.add_paragraph(f"{i}. {q_data['stem']}")
        
        # Options
        for j, option in enumerate(q_data['options']):
            label = chr(ord('A') + j)
            doc.add_paragraph(f"{label}. {option}")
        
        # Answer
        if q_data['question_type'] == 'true_false':
            answer_text = "True" if q_data['correct_answer'] == [0] else "False"
        else:
            answer_labels = [chr(ord('A') + idx) for idx in q_data['correct_answer']]
            answer_text = ", ".join(answer_labels)
        
        doc.add_paragraph(f"Answer: {answer_text}")
        
        # Explanation (if provided)
        if q_data.get('explanation'):
            doc.add_paragraph(f"Explanation: {q_data['explanation']}")
        
        # Add spacing between questions
        doc.add_paragraph("")
    
    doc.save(file_path)