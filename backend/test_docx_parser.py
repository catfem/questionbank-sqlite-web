import pytest
import tempfile
import os
import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from docx_parser import DocxParser, ParsedQuestion, create_sample_docx

class TestDocxParser:
    
    def test_single_choice_question(self):
        """Test parsing a single choice question."""
        questions_data = [
            {
                'stem': 'What is the capital of France?',
                'options': ['London', 'Berlin', 'Paris', 'Madrid'],
                'correct_answer': [2],
                'question_type': 'single',
                'explanation': 'Paris is the capital of France.'
            }
        ]
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            create_sample_docx(tmp.name, questions_data)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 1
            assert len(errors) == 0
            
            q = parsed_questions[0]
            assert 'capital of France' in q.stem
            assert q.question_type == 'single'
            assert len(q.options) == 4
            assert q.correct_answer == [2]
            assert 'Paris' in q.explanation
            
        os.unlink(tmp.name)
    
    def test_multiple_choice_question(self):
        """Test parsing a multiple choice question."""
        # Create a custom document for multiple choice
        doc = Document()
        doc.add_paragraph("1. Which of the following are programming languages?")
        doc.add_paragraph("A. Python")
        doc.add_paragraph("B. HTML")
        doc.add_paragraph("C. Java")
        doc.add_paragraph("D. CSS")
        doc.add_paragraph("Answer: A, C")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 1
            assert len(errors) == 0
            
            q = parsed_questions[0]
            assert q.question_type == 'multiple'
            assert len(q.options) == 4
            assert set(q.correct_answer) == {0, 2}  # A and C
            
        os.unlink(tmp.name)
    
    def test_true_false_question(self):
        """Test parsing a true/false question."""
        doc = Document()
        doc.add_paragraph("1. The Earth is flat.")
        doc.add_paragraph("Answer: False")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 1
            assert len(errors) == 0
            
            q = parsed_questions[0]
            assert q.question_type == 'true_false'
            assert len(q.options) == 2
            assert q.options[0]['text'] == 'True'
            assert q.options[1]['text'] == 'False'
            assert q.correct_answer == [1]  # False
            
        os.unlink(tmp.name)
    
    def test_missing_option_markers(self):
        """Test handling of questions with missing option markers."""
        doc = Document()
        doc.add_paragraph("1. What is 2 + 2?")
        doc.add_paragraph("The answer is four")
        doc.add_paragraph("A. 3")
        doc.add_paragraph("B. 4")
        doc.add_paragraph("Answer: B")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            # Should parse successfully despite extra text
            assert len(parsed_questions) == 1
            q = parsed_questions[0]
            assert '2 + 2' in q.stem
            assert len(q.options) == 2
            
        os.unlink(tmp.name)
    
    def test_multiline_question_stem(self):
        """Test parsing questions with multi-line stems."""
        doc = Document()
        doc.add_paragraph("1. Consider the following scenario:")
        doc.add_paragraph("A company has 100 employees.")
        doc.add_paragraph("If 20% work remotely, how many work in the office?")
        doc.add_paragraph("A. 20")
        doc.add_paragraph("B. 80")
        doc.add_paragraph("C. 100")
        doc.add_paragraph("D. 0")
        doc.add_paragraph("Answer: B")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 1
            assert len(errors) == 0
            
            q = parsed_questions[0]
            assert 'company has 100 employees' in q.stem
            assert 'work in the office' in q.stem
            assert q.correct_answer == [1]  # B
            
        os.unlink(tmp.name)
    
    def test_answer_on_next_paragraph(self):
        """Test parsing when answer is on a separate paragraph."""
        doc = Document()
        doc.add_paragraph("1. What is the largest planet in our solar system?")
        doc.add_paragraph("A. Mars")
        doc.add_paragraph("B. Jupiter")
        doc.add_paragraph("C. Saturn")
        doc.add_paragraph("D. Earth")
        doc.add_paragraph("")
        doc.add_paragraph("Answer: B")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 1
            assert len(errors) == 0
            
            q = parsed_questions[0]
            assert 'largest planet' in q.stem
            assert q.correct_answer == [1]  # B
            
        os.unlink(tmp.name)
    
    def test_error_handling_missing_answer(self):
        """Test error handling for questions without answers."""
        doc = Document()
        doc.add_paragraph("1. What color is the sky?")
        doc.add_paragraph("A. Red")
        doc.add_paragraph("B. Blue")
        doc.add_paragraph("C. Green")
        # Missing answer line
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 0
            assert len(errors) == 1
            assert 'Missing correct answer' in errors[0]['error']
            
        os.unlink(tmp.name)
    
    def test_error_handling_missing_options(self):
        """Test error handling for questions without options."""
        doc = Document()
        doc.add_paragraph("1. What is the meaning of life?")
        doc.add_paragraph("Answer: 42")
        # Missing options
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 0
            assert len(errors) == 1
            assert 'Missing options' in errors[0]['error']
            
        os.unlink(tmp.name)
    
    def test_different_question_number_formats(self):
        """Test different question number formats."""
        doc = Document()
        doc.add_paragraph("1) First question format")
        doc.add_paragraph("A. Option A")
        doc.add_paragraph("B. Option B")
        doc.add_paragraph("Answer: A")
        
        doc.add_paragraph("")
        doc.add_paragraph("Question 2: Second question format")
        doc.add_paragraph("A. Option A")
        doc.add_paragraph("B. Option B")
        doc.add_paragraph("Answer: B")
        
        doc.add_paragraph("")
        doc.add_paragraph("Q3. Third question format")
        doc.add_paragraph("A. Option A")
        doc.add_paragraph("B. Option B")
        doc.add_paragraph("Answer: A")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            parser = DocxParser()
            parsed_questions, errors = parser.parse_document(tmp.name)
            
            assert len(parsed_questions) == 3
            assert len(errors) == 0
            
        os.unlink(tmp.name)